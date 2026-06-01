"""Generate the final technical report for the eye-tracking deception project."""

from __future__ import annotations

import csv
import shutil
import subprocess
from datetime import date
from pathlib import Path

try:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt
except ImportError as exc:  # pragma: no cover - import guard for local environments
    raise SystemExit(
        "python-docx is required to generate the Word report. "
        "Install it with: pip install python-docx"
    ) from exc


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = PROJECT_ROOT / "reports" / "final_report"
DOCX_PATH = OUTPUT_DIR / "eye_tracking_deception_report.docx"
PDF_PATH = OUTPUT_DIR / "eye_tracking_deception_report.pdf"


def read_text(relative_path: str) -> str:
    path = PROJECT_ROOT / relative_path
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="replace")


def read_csv_rows(relative_path: str) -> list[dict[str, str]]:
    path = PROJECT_ROOT / relative_path
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_table(document: Document, caption: str, headers: list[str], rows: list[list[object]]) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(caption)
    run.bold = True
    run.italic = True

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], value)

    document.add_paragraph()


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.line_spacing = 1.15
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_heading(document: Document, text: str, level: int = 1, page_break: bool = False) -> None:
    if page_break and len(document.paragraphs) > 0:
        document.add_page_break()
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"


def add_toc_placeholder(document: Document) -> None:
    add_heading(document, "Table of Contents", level=1, page_break=True)
    add_paragraph(
        document,
        "This page is a table of contents placeholder. In Microsoft Word, right-click this "
        "placeholder area and update fields if automatic TOC generation is required. "
        "The report itself uses Word heading styles for all numbered sections.",
    )
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run._r.append(fld_char3)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(20)
    run = paragraph.add_run("Real-Time Eye-Tracking-Based Deception Risk Detection Using Neural Networks")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(18)

    for line in [
        "Technical and Research Report",
        "Prepared by: Alisher Temirkhan",
        "Astana IT University / Nazarbayev University",
        f"Date: {date.today().strftime('%B %d, %Y')}",
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(line)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    document.add_paragraph()
    add_paragraph(
        document,
        "This report documents a completed research prototype and planned full-system "
        "extension for controlled eye-tracking-based deception risk estimation. The system "
        "outputs calibrated risk scores and must not be interpreted as a universal lie "
        "detector or as an autonomous decision-making tool.",
    )


def metric(value: str) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return value or "N/A"


def prototype_summary_rows() -> list[list[str]]:
    preprocessing = read_text("data/processed/preprocessing_report.txt")
    split = read_text("data/processed/split_report.txt")
    baseline = read_text("reports/baselines/baseline_report.txt")

    def extract(prefix: str, text: str) -> str:
        for line in text.splitlines():
            if line.strip().startswith(prefix):
                return line.split(":", 1)[1].strip() if ":" in line else line.strip()
        return "Not available"

    return [
        ["Participants in preprocessing report", extract("participants:", preprocessing)],
        ["Sessions", extract("sessions:", preprocessing)],
        ["Trials", extract("trials:", preprocessing)],
        ["Raw gaze samples", extract("raw_gaze_samples:", preprocessing)],
        ["Windows created", extract("windows_created:", preprocessing)],
        ["Usable windows", extract("usable_windows:", preprocessing)],
        ["Train samples", extract("Train samples:", baseline)],
        ["Validation samples", extract("Validation samples:", baseline)],
        ["Test samples", extract("Test samples:", baseline)],
        ["Feature count", extract("Number of features used:", baseline)],
        ["Leakage check", "OK" if "Leakage check: OK" in split else "Not available"],
    ]


def current_metric_rows() -> list[list[str]]:
    rows = read_csv_rows("reports/baselines/baseline_metrics.csv")
    if not rows:
        return [["No baseline metric file was found.", "", "", "", "", "", "", "", "", ""]]
    output = []
    for row in rows:
        output.append(
            [
                row.get("split", ""),
                row.get("model_name", ""),
                metric(row.get("accuracy", "")),
                metric(row.get("balanced_accuracy", "")),
                metric(row.get("precision", "")),
                metric(row.get("recall", "")),
                metric(row.get("f1", "")),
                metric(row.get("roc_auc", "")),
                metric(row.get("false_positive_rate", "")),
                metric(row.get("false_negative_rate", "")),
            ]
        )
    return output


def add_project_sections(document: Document) -> None:
    preprocessing_report = read_text("data/processed/preprocessing_report.txt")
    split_report = read_text("data/processed/split_report.txt")
    baseline_report = read_text("reports/baselines/baseline_report.txt")

    add_heading(document, "1. Abstract", page_break=True)
    add_paragraph(
        document,
        "Deception detection is a scientifically difficult problem because behavioral cues "
        "are weak, context-dependent, and influenced by participant strategy, cognitive "
        "load, anxiety, fatigue, device quality, and experimental design. This project "
        "therefore treats eye tracking as a behavioral signal-analysis problem rather than "
        "as a direct truth-verification method. The prototype described in this report "
        "uses controlled statement-level tasks in which participants are instructed to "
        "answer truthfully or deceptively while gaze position, pupil diameter, blink, "
        "fixation, saccade, and validity signals are recorded.",
    )
    add_paragraph(
        document,
        "The system is designed to estimate a probability score for deception-related risk "
        "patterns under controlled experimental conditions. It does not prove whether a "
        "person is lying and must not be used as a legal, employment, medical, or "
        "security-critical decision tool. The current implementation includes project "
        "definition, experiment protocol documentation, a Tkinter-based experiment "
        "interface, mock eye-tracker data collection, raw CSV storage, preprocessing, "
        "sliding-window generation, subject-independent splitting, and baseline machine "
        "learning models. The planned full version extends this pipeline with LSTM, GRU, "
        "and causal Temporal Convolutional Network models, a real-time inference service, "
        "and a monitoring dashboard.",
    )

    add_heading(document, "2. Introduction", page_break=True)
    add_paragraph(
        document,
        "Deception-related behavior is not expressed through a single reliable cue. A "
        "controlled participant may maintain eye contact, deliberately suppress visible "
        "reactions, or use countermeasures that change gaze behavior. For this reason, "
        "modern computational approaches should not be framed as universal lie detection. "
        "A more defensible framing is to estimate risk patterns that are statistically "
        "associated with instructed deceptive responses in a defined protocol.",
    )
    add_paragraph(
        document,
        "Eye tracking provides continuous, high-frequency signals that can reflect visual "
        "attention, search behavior, cognitive effort, hesitation, and oculomotor changes. "
        "Pupil diameter may vary with cognitive load, blink and fixation patterns may "
        "reflect attention allocation, and saccade dynamics may reveal changes in visual "
        "processing strategy. Neural sequence models are relevant because deception-related "
        "patterns may unfold over time rather than appear in a single static summary. "
        "However, real-time deployment adds latency, missing-data, calibration, and causal "
        "preprocessing constraints.",
    )

    add_heading(document, "3. Research Motivation", page_break=True)
    add_paragraph(
        document,
        "The motivation for this project is to create a reproducible prototype pipeline "
        "that connects controlled experimental design with machine learning evaluation. "
        "Behavioral cues such as gaze dispersion, fixation duration, pupil variability, "
        "and blink timing are not direct evidence of lying. They are indirect signals that "
        "may correlate with attention, inhibition, monitoring, and cognitive control. A "
        "scientific prototype must therefore preserve trial labels, participant identity, "
        "signal quality indicators, and evaluation splits so that model performance can be "
        "tested without participant leakage.",
    )
    add_paragraph(
        document,
        "The article basis for this project, A Systematic Review of Neural Network "
        "Analysis of Real-Time Eye-Tracking Data for Detecting Deceptive Statements, "
        "motivates an emphasis on causal preprocessing, subject-independent validation, "
        "real-time feasibility, and cautious interpretation. These requirements are "
        "reflected in the repository design and in this report.",
    )

    add_heading(document, "4. Problem Statement", page_break=True)
    add_paragraph(
        document,
        "At the prototype level, the task is binary supervised classification. Each "
        "sliding window belongs to a trial with a label derived from the instruction: "
        "0 for truthful response and 1 for deceptive response. The model input includes "
        "raw and derived eye-tracking variables such as gaze_x, gaze_y, pupil_left, "
        "pupil_right, blink, fixation, saccade, validity, pupil_mean, gaze_velocity, "
        "missing-data features, and aggregated window statistics.",
    )
    add_paragraph(
        document,
        "The output is a deception risk score between 0 and 1. The interface may map that "
        "score into low, medium, or high risk for decision support, but the score is not a "
        "final truth or lie judgment. When signal quality is poor, the system should return "
        "insufficient data instead of forcing a classification.",
    )
    add_table(
        document,
        "Table 1. Prototype input and output definition.",
        ["Element", "Definition"],
        [
            ["Input signal", "gaze_x, gaze_y, pupil_left, pupil_right, blink, fixation, saccade, validity"],
            ["Derived features", "pupil_mean, gaze_velocity, missing_ratio, valid_ratio, window statistics"],
            ["Prototype label", "0 = truthful instruction, 1 = deceptive instruction"],
            ["Model output", "Calibrated deception risk score from 0 to 1"],
            ["Operational status", "low risk, medium risk, high risk, or insufficient data"],
        ],
    )

    add_heading(document, "5. Scope and Limitations", page_break=True)
    add_paragraph(
        document,
        "The system is scoped to controlled experimental conditions. It is not a universal "
        "lie detector, it does not infer moral truthfulness, and it is not appropriate for "
        "autonomous legal, employment, medical, or security-critical decisions. Eye-tracking "
        "signals are sensitive to illumination, screen geometry, device calibration, glasses, "
        "fatigue, participant variability, question content, and deliberate countermeasures. "
        "Any deployable version would require informed consent, careful study design, "
        "participant-level validation, and transparent uncertainty reporting.",
    )

    add_heading(document, "6. Literature and Scientific Basis", page_break=True)
    add_paragraph(
        document,
        "Relevant research includes statement-level deception tasks, concealed information "
        "tests, concealed recognition paradigms, and machine learning analysis of oculomotor "
        "signals. Eye-tracking features commonly include gaze position, dwell time, fixation "
        "count, fixation duration, saccade rate, saccade amplitude, pupil diameter, blink "
        "rate, and missing-data patterns. Neural sequence models such as LSTM and GRU "
        "networks can model temporal dependencies, while causal TCN architectures are "
        "attractive for low-latency real-time processing.",
    )
    add_paragraph(
        document,
        "Methodological risks are central in this domain. Randomly splitting windows can "
        "cause the same participant's data to appear in both training and test sets, "
        "inflating performance estimates. Non-causal preprocessing can leak future samples "
        "into the current decision. Missing-data handling can accidentally encode device "
        "or participant identity. The prototype design therefore emphasizes participant "
        "splits, leakage checks, validity ratios, and explicit reporting of false positives "
        "and false negatives.",
    )

    add_heading(document, "7. System Overview", page_break=True)
    add_paragraph(
        document,
        "The full system architecture connects an experiment interface, eye-tracking data "
        "stream, raw CSV writer, preprocessing module, window generator, participant-level "
        "splitter, baseline model trainer, planned neural sequence models, planned real-time "
        "inference service, and planned dashboard. The repository currently implements the "
        "core offline pipeline up to baseline model evaluation.",
    )
    add_table(
        document,
        "Table 2. System components and implementation status.",
        ["Component", "Responsibility", "Current status"],
        [
            ["Project definition", "Defines task, outputs, and scientific limits", "Implemented"],
            ["Experiment protocol", "Defines controlled truth, lie, and mixed blocks", "Implemented"],
            ["Tkinter interface", "Runs local statement-level task", "Prototype implemented"],
            ["Mock eye tracker", "Generates prototype gaze samples without hardware", "Prototype implemented"],
            ["Raw CSV writer", "Stores participants, sessions, trials, and gaze samples", "Implemented"],
            ["Preprocessing", "Validates, cleans, derives features, reports quality", "Implemented"],
            ["Sliding windows", "Creates trial-level 3-second windows with 1-second stride", "Implemented"],
            ["Subject split", "Splits by participant and checks leakage", "Implemented"],
            ["Baseline ML", "Trains Logistic Regression and Random Forest", "Implemented"],
            ["Neural sequence models", "LSTM, GRU, and causal TCN sequence classifiers", "Planned full version"],
            ["Real-time inference", "Streaming buffer, causal features, risk scores", "Planned full version"],
            ["Dashboard", "Displays risk score, signal quality, and system state", "Planned full version"],
        ],
    )

    add_heading(document, "8. Experimental Protocol", page_break=True)
    add_paragraph(
        document,
        "The protocol follows a statement-level deception task. A participant registers "
        "with an anonymous identifier, completes eye-tracker calibration, views a baseline "
        "fixation point, and then answers questions under truthful or deceptive instructions. "
        "The key label is based on the instruction shown to the participant, not merely on "
        "the yes/no answer. This matters because the same answer can be truthful or "
        "deceptive depending on the participant's actual state and the instruction.",
    )
    add_table(
        document,
        "Table 3. Trial and block structure.",
        ["Stage", "Duration or count", "Purpose"],
        [
            ["Participant registration", "Before task", "Assign anonymous participant identifier"],
            ["Calibration", "Before task", "Estimate eye-tracker quality and geometry"],
            ["Baseline recording", "60 seconds", "Record neutral fixation for normalization"],
            ["Fixation cross", "2 seconds", "Stabilize attention before each question"],
            ["Instruction screen", "1 second", "Tell participant to answer truthfully or deceptively"],
            ["Question screen", "Up to 5 seconds", "Collect response and gaze behavior"],
            ["Pause", "1 second", "Separate consecutive trials"],
            ["Truth block", "10 trials", "Label 0"],
            ["Lie block", "10 trials", "Label 1"],
            ["Mixed block", "20 trials", "Randomized truth and lie instructions"],
        ],
    )

    add_heading(document, "9. Dataset Design", page_break=True)
    add_paragraph(
        document,
        "The raw dataset is normalized into four CSV files. This structure preserves "
        "participant-level metadata, session metadata, trial labels, and high-frequency "
        "gaze samples while avoiding storage of personal names. The relational structure "
        "is participant to session, session to trial, and trial to gaze_samples.",
    )
    add_table(
        document,
        "Table 4. Raw dataset file responsibilities.",
        ["File", "Role", "Primary relationship"],
        [
            ["participants.csv", "Anonymous participant records", "participant_id"],
            ["sessions.csv", "Recording session metadata", "participant_id -> session_id"],
            ["trials.csv", "Question, instruction, response, and label", "session_id -> trial_id"],
            ["gaze_samples.csv", "Time-indexed raw eye-tracking samples", "trial_id -> sample_id"],
        ],
    )
    add_table(
        document,
        "Table 5. participants.csv schema.",
        ["Column", "Type", "Description"],
        [["participant_id", "string", "Anonymous participant identifier"], ["notes", "string", "Non-identifying notes"]],
    )
    add_table(
        document,
        "Table 6. sessions.csv schema.",
        ["Column", "Type", "Description"],
        [
            ["session_id", "string", "Unique recording session identifier"],
            ["participant_id", "string", "Participant who completed the session"],
            ["date", "date/string", "Recording date"],
            ["device", "string", "Eye tracker or mock device"],
            ["screen_width", "integer", "Screen width in pixels"],
            ["screen_height", "integer", "Screen height in pixels"],
            ["sampling_rate", "numeric", "Eye-tracking sampling rate in Hz"],
            ["calibration_quality", "string", "Calibration quality label"],
        ],
    )
    add_table(
        document,
        "Table 7. trials.csv schema.",
        ["Column", "Type", "Description"],
        [
            ["trial_id", "string", "Unique trial identifier"],
            ["session_id", "string", "Session containing the trial"],
            ["question_text", "string", "Displayed question"],
            ["instruction", "string", "truth or lie instruction"],
            ["label", "integer", "0 for truth, 1 for lie"],
            ["answer", "string", "Participant yes/no answer"],
            ["response_time", "numeric", "Time to answer in seconds"],
            ["start_time", "numeric", "Trial start time"],
            ["end_time", "numeric", "Trial end time"],
        ],
    )
    add_table(
        document,
        "Table 8. gaze_samples.csv schema.",
        ["Column", "Type", "Description"],
        [
            ["sample_id", "integer", "Sample identifier"],
            ["trial_id", "string", "Trial containing the sample"],
            ["timestamp", "numeric", "Time inside the trial"],
            ["gaze_x", "numeric", "Normalized horizontal gaze position"],
            ["gaze_y", "numeric", "Normalized vertical gaze position"],
            ["pupil_left", "numeric", "Left pupil diameter"],
            ["pupil_right", "numeric", "Right pupil diameter"],
            ["blink", "binary", "Blink indicator"],
            ["fixation", "binary", "Fixation indicator"],
            ["saccade", "binary", "Saccade indicator"],
            ["validity", "binary", "1 when the sample is valid"],
        ],
    )

    add_heading(document, "10. Data Collection Interface", page_break=True)
    add_paragraph(
        document,
        "The data collection prototype uses a local Tkinter application. It provides a "
        "start screen, participant input, calibration placeholder, baseline placeholder, "
        "trial loop, yes/no response controls, and CSV output. The current prototype uses "
        "a mock eye tracker, which allows the full software pipeline to be developed and "
        "tested before physical eye-tracking hardware is integrated.",
    )
    add_paragraph(
        document,
        "Using a mock tracker is scientifically limited because it cannot validate true "
        "oculomotor behavior. Its role is engineering validation: it verifies that session "
        "metadata, trial labels, timestamps, sample rows, preprocessing, split logic, and "
        "training scripts can operate end to end.",
    )

    add_heading(document, "11. Preprocessing Pipeline", page_break=True)
    add_paragraph(
        document,
        "Preprocessing validates raw CSV columns, checks relational integrity, verifies "
        "label consistency, validates gaze coordinate ranges, validates pupil values, "
        "handles missing or invalid samples, computes validity ratios, derives pupil_mean, "
        "and estimates gaze_velocity. For real-time use, preprocessing must be causal: a "
        "decision at time t may use current and past samples only, not future samples.",
    )
    add_paragraph(
        document,
        "Pipeline diagram: Raw gaze samples -> validation -> cleaning -> feature calculation "
        "-> sliding windows -> processed dataset."
    )
    add_table(
        document,
        "Table 9. Preprocessing checks and derived variables.",
        ["Step", "Description"],
        [
            ["CSV validation", "Check required columns and readable files"],
            ["Relational checks", "Verify participant, session, trial, and sample links"],
            ["Label checks", "Confirm truth maps to 0 and lie maps to 1"],
            ["Signal checks", "Validate gaze coordinates, pupil values, and binary indicators"],
            ["Quality metrics", "Calculate valid_ratio and missing_ratio"],
            ["Feature derivation", "Calculate pupil_mean and gaze_velocity"],
            ["Causal principle", "Avoid future-sample leakage for real-time decisions"],
        ],
    )

    add_heading(document, "12. Sliding Window Generation", page_break=True)
    add_paragraph(
        document,
        "The implemented windowing stage uses 3-second windows with 1-second stride. Each "
        "window belongs to one trial and inherits the trial label. A window is marked usable "
        "when its valid sample ratio meets the minimum threshold. Aggregated features are "
        "computed for gaze position, pupil diameter, blink rate, fixation rate, saccade "
        "rate, velocity, and quality indicators.",
    )
    add_table(
        document,
        "Table 10. Sliding-window feature groups.",
        ["Feature group", "Examples"],
        [
            ["Gaze position", "gaze_x_mean, gaze_x_std, gaze_y_mean, gaze_y_std"],
            ["Pupil", "pupil_left_mean, pupil_right_mean, pupil_mean_mean, pupil_mean_std"],
            ["Blink", "blink_rate"],
            ["Fixation and saccade", "fixation_rate, saccade_rate"],
            ["Velocity", "gaze_velocity_mean, gaze_velocity_std"],
            ["Quality", "valid_ratio, missing_ratio, is_usable"],
        ],
    )

    add_heading(document, "13. Subject-Independent Split", page_break=True)
    add_paragraph(
        document,
        "A random window split is inappropriate because windows from the same participant "
        "share physiology, calibration, screen setup, and task strategy. If windows from "
        "one participant appear in both training and test sets, evaluation can measure "
        "participant recognition rather than deception-related generalization. The project "
        "therefore uses participant-level splitting: no participant_id may appear in more "
        "than one split.",
    )
    add_table(
        document,
        "Table 11. Implemented split policy.",
        ["Property", "Value"],
        [
            ["Split unit", "participant_id"],
            ["Train ratio", "0.70"],
            ["Validation ratio", "0.15"],
            ["Test ratio", "0.15"],
            ["Random seed", "42"],
            ["Use only usable windows", "True"],
            ["Leakage rule", "No participant_id in more than one split"],
        ],
    )

    add_heading(document, "14. Baseline Models", page_break=True)
    add_paragraph(
        document,
        "The implemented baseline models are Logistic Regression and Random Forest. Logistic "
        "Regression uses standardized features and provides an interpretable linear baseline. "
        "Random Forest provides a nonlinear tabular-feature baseline. Class weighting is used "
        "to reduce the effect of class imbalance when present. Baselines are necessary before "
        "neural networks because they establish whether engineered features contain a usable "
        "signal and whether the data pipeline is healthy.",
    )
    add_table(
        document,
        "Table 12. Baseline model roles.",
        ["Model", "Role", "Strength", "Limitation"],
        [
            ["Logistic Regression", "Linear baseline", "Simple, calibrated, interpretable", "Cannot model nonlinear temporal dynamics"],
            ["Random Forest", "Nonlinear feature baseline", "Handles interactions and thresholds", "Uses aggregated windows, not raw sequence dynamics"],
        ],
    )

    add_heading(document, "15. Neural Sequence Models", page_break=True)
    add_paragraph(
        document,
        "The full prototype is planned to include LSTM, GRU, and causal TCN models. Their "
        "input shape is batch_size x time_steps x features. LSTM and GRU networks process "
        "temporal dependencies through gated recurrent units. A causal TCN uses one-dimensional "
        "convolutions with temporal padding that prevents future leakage, making it suitable "
        "for low-latency streaming inference.",
    )
    add_table(
        document,
        "Table 13. Planned neural sequence model comparison.",
        ["Model", "Temporal mechanism", "Real-time suitability", "Expected use"],
        [
            ["LSTM", "Input, forget, and output gates", "Moderate latency", "Model longer eye-movement dependencies"],
            ["GRU", "Reset and update gates", "Moderate to low latency", "Compact recurrent sequence baseline"],
            ["Causal TCN", "Causal dilated convolutions", "High suitability", "Low-latency window inference"],
        ],
    )

    add_heading(document, "16. Real-Time Inference Architecture", page_break=True)
    add_paragraph(
        document,
        "The planned real-time architecture receives samples from an eye tracker, stores them "
        "in a rolling buffer, applies causal preprocessing, creates sliding windows, runs "
        "model inference, smooths consecutive risk estimates, and displays a calibrated "
        "risk score with signal-quality status. The dashboard should display insufficient "
        "data when the valid sample ratio is below the configured threshold.",
    )
    add_table(
        document,
        "Table 14. Decision-support policy.",
        ["Condition", "Output"],
        [
            ["valid_ratio < 0.70", "Insufficient data"],
            ["score < 0.40", "Low deception risk"],
            ["0.40 <= score <= 0.69", "Medium deception risk"],
            ["score >= 0.70", "High deception risk"],
        ],
    )

    add_heading(document, "17. Evaluation Methodology", page_break=True)
    add_paragraph(
        document,
        "Evaluation should report accuracy, balanced accuracy, precision, recall, F1, ROC-AUC, "
        "average precision, false positive rate, false negative rate, mean inference latency, "
        "valid sample ratio, and participant leakage checks. Balanced accuracy is important "
        "when classes are imbalanced. False positive rate is especially important because "
        "wrongly flagging a truthful participant as high risk can cause serious harm if "
        "the system is misused.",
    )
    add_table(
        document,
        "Table 15. Evaluation metrics.",
        ["Metric", "Purpose"],
        [
            ["Accuracy", "Overall fraction of correct labels"],
            ["Balanced accuracy", "Mean recall across classes"],
            ["Precision", "Reliability of positive risk predictions"],
            ["Recall", "Sensitivity to deceptive-instruction windows"],
            ["F1", "Harmonic mean of precision and recall"],
            ["ROC-AUC", "Ranking quality across thresholds"],
            ["Average precision", "Precision-recall performance"],
            ["False positive rate", "Truth windows incorrectly flagged as lie risk"],
            ["False negative rate", "Lie windows missed by the system"],
            ["Latency", "Real-time feasibility per window"],
        ],
    )

    add_heading(document, "18. Current Prototype Results", page_break=True)
    add_paragraph(
        document,
        "This section uses repository-generated artifacts when available. These values are "
        "current prototype results from the local pipeline, not large-scale human-subject "
        "evidence. The raw data in this repository is sufficient to exercise preprocessing, "
        "splitting, and baseline training, but it remains too small for strong scientific "
        "claims about generalization.",
    )
    add_table(document, "Table 16. Current local prototype summary.", ["Item", "Value"], prototype_summary_rows())
    add_table(
        document,
        "Table 17. Current local baseline metrics from reports/baselines/baseline_metrics.csv.",
        ["Split", "Model", "Accuracy", "Balanced acc.", "Precision", "Recall", "F1", "ROC-AUC", "FPR", "FNR"],
        current_metric_rows(),
    )
    if preprocessing_report:
        add_paragraph(document, "Preprocessing report note: " + preprocessing_report.replace("\n", " ")[:1200])
    if split_report:
        add_paragraph(document, "Split report note: " + split_report.replace("\n", " ")[:1200])
    if baseline_report:
        add_paragraph(document, "Baseline report note: " + baseline_report.replace("\n", " ")[:1500])

    add_heading(document, "19. Synthetic Full-Scale Prototype Evaluation", page_break=True)
    add_paragraph(
        document,
        "The values in this section are illustrative synthetic results used to demonstrate "
        "the expected reporting format of the full evaluation. They must be replaced by "
        "real experimental values after collecting a sufficient participant dataset. They "
        "must not be reported as measured human-subject findings.",
    )
    add_table(
        document,
        "Table 18. Synthetic illustrative dataset summary.",
        ["Property", "Synthetic value"],
        [
            ["Participants", "36"],
            ["Sessions", "36"],
            ["Trials per participant", "40"],
            ["Total trials", "1,440"],
            ["Truth trials", "720"],
            ["Lie trials", "720"],
            ["Sampling rate", "60 Hz"],
            ["Window size", "3 seconds"],
            ["Window stride", "1 second"],
            ["Total generated windows", "4,320"],
            ["Usable windows after filtering", "4,018"],
            ["Minimum valid sample ratio", "0.70"],
        ],
    )
    add_table(
        document,
        "Table 19. Synthetic illustrative subject-independent split.",
        ["Split", "Participants", "Windows"],
        [["Train", "25", "2,790"], ["Validation", "5", "558"], ["Test", "6", "670"], ["Leakage check", "OK", "OK"]],
    )
    add_table(
        document,
        "Table 20. Synthetic illustrative baseline results.",
        ["Model", "Val acc.", "Val bal. acc.", "Val F1", "Val ROC-AUC", "Test acc.", "Test bal. acc.", "Test F1", "Test ROC-AUC", "FPR", "FNR"],
        [
            ["Logistic Regression", "0.64", "0.63", "0.62", "0.68", "0.62", "0.61", "0.60", "0.66", "0.34", "0.42"],
            ["Random Forest", "0.69", "0.68", "0.68", "0.74", "0.67", "0.66", "0.66", "0.72", "0.31", "0.36"],
        ],
    )
    add_table(
        document,
        "Table 21. Synthetic illustrative neural sequence model results.",
        ["Model", "Test acc.", "Bal. acc.", "Precision", "Recall", "F1", "ROC-AUC", "Avg. precision", "FPR", "FNR", "Latency"],
        [
            ["LSTM", "0.71", "0.70", "0.72", "0.68", "0.70", "0.77", "0.75", "0.28", "0.32", "24 ms"],
            ["GRU", "0.70", "0.69", "0.71", "0.67", "0.69", "0.76", "0.74", "0.29", "0.33", "21 ms"],
            ["Causal TCN", "0.75", "0.74", "0.76", "0.72", "0.74", "0.81", "0.79", "0.24", "0.28", "16 ms"],
        ],
    )
    add_table(
        document,
        "Table 22. Synthetic illustrative latency comparison.",
        ["Model", "Mean inference latency per window", "Real-time interpretation"],
        [["LSTM", "24 ms", "Feasible with buffering"], ["GRU", "21 ms", "Feasible with lower recurrent cost"], ["Causal TCN", "16 ms", "Most attractive for low-latency deployment"]],
    )
    add_table(
        document,
        "Table 23. Synthetic illustrative error comparison.",
        ["Model", "False positive rate", "False negative rate", "Interpretation"],
        [
            ["Logistic Regression", "0.34", "0.42", "Weak baseline with high missed-risk rate"],
            ["Random Forest", "0.31", "0.36", "Improved nonlinear baseline"],
            ["LSTM", "0.28", "0.32", "Temporal model improves both error types"],
            ["GRU", "0.29", "0.33", "Similar to LSTM with slightly lower latency"],
            ["Causal TCN", "0.24", "0.28", "Best synthetic tradeoff"],
        ],
    )
    add_paragraph(
        document,
        "The synthetic results suggest that temporal neural models may outperform aggregated "
        "feature baselines, with the causal TCN offering the strongest synthetic balance of "
        "ROC-AUC, F1, error rates, and latency. This is only a hypothesis until validated "
        "with real participants, cross-session testing, cross-device testing, and "
        "countermeasure-aware evaluation.",
    )

    add_heading(document, "20. Discussion", page_break=True)
    add_paragraph(
        document,
        "The current prototype demonstrates an end-to-end pipeline from controlled task "
        "definition to baseline model reports. Its most important scientific strength is "
        "not the current metric values, which come from a very small dataset, but the "
        "structure of the pipeline: explicit labels, documented raw schemas, validity "
        "checks, sliding windows, participant-level splitting, and cautious result wording. "
        "The synthetic evaluation illustrates why temporal sequence models are a plausible "
        "next direction, especially when deception-related behavior depends on patterns "
        "over several seconds rather than isolated summary statistics.",
    )
    add_paragraph(
        document,
        "The causal TCN is promising for real-time deployment because convolutional sequence "
        "processing can be parallelized and constrained to past samples. Nevertheless, gaze-only "
        "classification remains limited. Any observed signal may reflect attention, stress, "
        "question difficulty, memory, or fatigue rather than deception itself. Controlled "
        "protocols and careful controls are therefore essential.",
    )

    add_heading(document, "21. Ethical Considerations", page_break=True)
    add_paragraph(
        document,
        "Eye-tracking data can reveal sensitive information about attention, fatigue, visual "
        "processing, and cognitive state. Participants should provide informed consent, data "
        "should be anonymized, and access should be restricted. The system must not present "
        "a risk score as a final lie judgment. False positives can unfairly harm truthful "
        "participants, while false negatives can produce misplaced confidence. The prototype "
        "should be used only as a research and decision-support tool under controlled "
        "conditions.",
    )

    add_heading(document, "22. Limitations", page_break=True)
    add_paragraph(
        document,
        "The current implementation uses a mock eye tracker for prototype data collection, "
        "which limits behavioral validity. The local dataset is small, and the synthetic "
        "full-scale results are illustrative rather than measured. Real validation requires "
        "hardware integration, more participants, cross-session validation, cross-device "
        "testing, countermeasure-aware studies, illumination controls, and analysis of pupil "
        "measurement instability. The system may also be sensitive to glasses, screen size, "
        "tracking loss, participant fatigue, and question wording.",
    )

    add_heading(document, "23. Future Work", page_break=True)
    add_paragraph(
        document,
        "Future work should integrate a real eye tracker, collect a larger participant "
        "dataset, implement a real-time dashboard, train LSTM, GRU, and causal TCN models, "
        "calibrate probability outputs, test cross-session and cross-device generalization, "
        "and evaluate robustness against participant countermeasures. A multimodal extension "
        "could combine eye tracking with speech, facial video, response time, and task "
        "metadata, but such fusion would increase privacy and consent requirements.",
    )

    add_heading(document, "24. Conclusion", page_break=True)
    add_paragraph(
        document,
        "This project builds a structured prototype for real-time eye-tracking-based "
        "deception risk estimation. The current repository implements data collection, raw "
        "CSV storage, preprocessing, windowing, subject-independent splitting, and baseline "
        "models. The system estimates risk patterns under controlled experimental conditions; "
        "it does not determine truth and should not be used as an autonomous decision tool. "
        "Subject-independent validation, causal preprocessing, quality-aware inference, and "
        "transparent reporting are central to the scientific validity of the approach. "
        "Temporal neural models, especially causal TCNs, are the next strongest technical "
        "direction for the full prototype.",
    )

    add_heading(document, "25. References", page_break=True)
    references = [
        "Vrij et al. (2019). Reference cited in the project article on deception-related behavioral analysis.",
        "Celniak and Slapczynska (2023). Reference cited in the project article on eye-tracking and deception-related analysis.",
        "Cook et al. (2012). Reference cited in the project article on oculomotor indicators and deception-related tasks.",
        "Peth et al. (2013). Reference cited in the project article on concealed information and eye-tracking analysis.",
        "Lancry-Dayan et al. (2023). Reference cited in the project article on neural or computational analysis of eye-tracking data.",
        "Gallardo-Antolin and Montero (2021). Reference cited in the project article on neural sequence or behavioral signal modeling.",
        "Elmadjian et al. (2023). Reference cited in the project article on modern machine learning analysis for deception-related signals.",
        "Millen and Hancock (2019). Reference cited in the project article on deception, behavior, and computational analysis.",
    ]
    for reference in references:
        add_paragraph(document, reference)


def convert_to_pdf() -> bool:
    try:
        from docx2pdf import convert  # type: ignore

        convert(str(DOCX_PATH), str(PDF_PATH))
        return PDF_PATH.exists()
    except Exception as exc:
        print(f"PDF conversion with docx2pdf skipped: {exc}")

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        print("PDF conversion skipped: docx2pdf/Word or LibreOffice was not available.")
        return False

    try:
        subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                str(OUTPUT_DIR),
                str(DOCX_PATH),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        return PDF_PATH.exists()
    except Exception as exc:
        print(f"PDF conversion skipped: LibreOffice conversion failed: {exc}")
        return False


def build_report() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    add_title_page(document)
    add_toc_placeholder(document)
    add_project_sections(document)
    document.save(DOCX_PATH)

    pdf_created = convert_to_pdf()
    print(f"DOCX report generated: {DOCX_PATH}")
    if pdf_created:
        print(f"PDF report generated: {PDF_PATH}")
    else:
        print("PDF report was not generated; DOCX output is available.")


if __name__ == "__main__":
    build_report()
